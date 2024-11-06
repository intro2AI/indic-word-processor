from PyQt5 import QtCore,QtWidgets
import typing
from docx import Document
import os
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
import uuid
IMAGE_EXTENSIONS = ['svg','.jpg','.png','.bmp']
def hexuuid():
    return uuid.uuid4().hex

def splitext(p):
    return os.path.splitext(p)[1].lower()

class TextEdit(QTextEdit):
    def __init__(self,label_trackAllKeys, *args, **kwargs):
        super(self.__class__, self).__init__(*args, **kwargs)

        self.label_trackAllKeys = label_trackAllKeys
        self.current_script: str = ""
        # IMPORTING HINDI MAPPING DICTIONARIES
        from mappings.mappings_v1 import C,V,v,misc
        self.C = C #consonents
        self.V = V #Independent Vowels
        self.v = v #dependent vowels
        self.misc = misc #nukta, halant, numbers and misc symbols
        self.english_bypass = False
        self.is_backtick_pressed = False

        # Variables used for Language rules
        self.lastChars= '   '#str
        self.makeNextVowelDependent = False
        self.use2CharsVowelNext = True
    
    def pasteWithoutFormatting(self):
        clipboard = QtWidgets.QApplication.clipboard()
        text = clipboard.text()
        self.insertPlainText(text)
        font = self.document().defaultFont()
        self.document().setDefaultFont(font)


    def keyPressEvent(self, event):
        print("_______________________________")
        print("use2CharVowelNext = "+str(self.use2CharsVowelNext))
        print(f"keyPressed: {event.text()}")
        print(f"keyPressed_id: {event}")
        # Process the default keys (spacebar, backspace, return/enter) like how we would normally
        if event.key() in [QtCore.Qt.Key_Enter,QtCore.Qt.Key_Return]:
            self.makeNextVowelDependent = False # new line should start with an independent vowele
            self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'\n')
            #self.label_trackScriptKeys.setText(self.label_trackScriptKeys.text()+"\n")
            return super(TextEdit, self).keyPressEvent(event) 
        elif event.key() in [QtCore.Qt.Key_Left,QtCore.Qt.Key_Right]: #, QtCore.Qt.AltModifier
            return super(TextEdit, self).keyPressEvent(event) 
        elif event.key() in [QtCore.Qt.Key_Shift]:
            return super(TextEdit, self).keyPressEvent(event)
        elif event.key() in [QtCore.Qt.Key_F5]: # refresh all flags and clear context
                self.label_trackAllKeys.setText('   ')#str
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True
                print("F5 pressed")
                return super(TextEdit, self).keyPressEvent(event)
        # if int(event.modifiers()) == (QtCore.Qt.ControlModifier+QtCore.Qt.AltModifier):
        #         # Alt Gr or Ctrl+Alt pressed 
        #         print('Alt Gr or Ctrl+Alt pressed')
        #         #event.ignore()
        #         # THE FOLLOWING IS DEDICATED LOGIC FOR DEPENDENT SINDHI VOWELS āū,ōū,ō,ē
        #         # FOR THE INDEPENDENT VOWELS, SHIFT KEY IS ALSO PRESSED, HENCE THIS SECTION DOES NOT GET EXECUTED THEN.
                
        #         # TODO if key pressed is ōūā, then insert text(in display and in context), 
                # without that wierd alt functionality.

        if event.modifiers() & QtCore.Qt.ControlModifier:
            if event.key() == QtCore.Qt.Key_Z:
                self.undo()
            elif event.key() == QtCore.Qt.Key_C:
                self.copy()
            elif event.key() == QtCore.Qt.Key_V:
                self.pasteWithoutFormatting()
            elif event.key() == QtCore.Qt.Key_X:
                self.cut()
        elif event.modifiers() & QtCore.Qt.AltModifier:
            event.ignore()
            print('ALT PRESSED')
        elif event.key() == QtCore.Qt.Key_Backspace:
            if len(self.toPlainText()) >1: 
                # do backspace in the keypress tracker also..
                self.label_trackAllKeys.setText(self.label_trackAllKeys.text()[:-1])
                return super(TextEdit, self).keyPressEvent(event)
            else:
                print("CLEARED")
                self.clear()
        elif event.key() == QtCore.Qt.Key_Down:
            cursor = self.textCursor()
            cursor.movePosition(cursor.Start)
            self.setTextCursor(cursor)
            return super(TextEdit, self).keyPressEvent(event)
        elif event.key() == QtCore.Qt.Key_Up:
            cursor = self.textCursor()
            cursor.movePosition(cursor.End)
            self.setTextCursor(cursor)
            return super(TextEdit, self).keyPressEvent(event)
        

        elif self.english_bypass == True:
            print("in here typing ENLGISH")
            self.insertPlainText(event.text())
        elif self.is_backtick_pressed: #################
            if QtCore.Qt.Key_A <= event.key() <= QtCore.Qt.Key_Z:
                print(f"BACKTICK + {event.text().upper()} PRESSED")
                print("in here typing ENLGISH")
                self.insertPlainText(event.text())
                self.label_trackAllKeys.setText('   ')#str
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True
        elif event.key() == QtCore.Qt.Key_QuoteLeft:  # Backtick key
            self.is_backtick_pressed = True
        else:
            self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+event.text())
            self.lastChars = self.label_trackAllKeys.text()[-5:]

            cursor = self.textCursor()

            key_5 = ''.join(self.lastChars[-5:])
            key_4 = ''.join(self.lastChars[-4:])
            key_3 = ''.join(self.lastChars[-3:])
            key_2 = ''.join(self.lastChars[-2:])
            key_1 = ''.join(self.lastChars[-1:])

            # Handling errors at the start of the code
            while len(key_5)<5:
                key_5 = ' '+ key_5
            while len(key_4)<4:
                key_4 = ' '+ key_4
            while len(key_3)<3:
                key_3 = ' '+ key_3
            while len(key_2)<2:
                key_2 = ' '+ key_2

            print("5 keys are:")
            print(key_5)
            print(key_4)
            print(key_3)
            print(key_2)
            print(key_1)
        
            #EDGE CASES
            if key_5[0] in self.C and key_5[2] == 'r' and key_5[4] == 'u':
                    self.deletePreviousChars(4 ,cursor)
                    self.insertPlainText(self.C[key_5[0]]+self.v['Rri'])
        
            #CONSONENTS AND MISC
            #detect matching keys in the last 4,3,2,1 chars
            #independent when last is " ", or last is Dependent
            #Dependent when last is consonent, or last is in a,e,i,o,u
            elif key_5 in self.C:
                #delete -1, and then add new consonent
                print("in key_5")
                if key_5 in ['sqhqr']:
                    self.deletePreviousChars(2 ,cursor)
                    self.insertPlainText(self.C[key_5])
                    
                else:
                    self.deletePreviousChars(4 ,cursor)
                    self.insertPlainText(self.C[key_5])
                    if key_5 not in ['DqDqA']:
                        # AUTO ADD HALANT
                        self.insertPlainText(self.misc['q'])
                        self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q') #logging in english

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_4 in self.C:
                print("in key_4")
                if key_4 in ['Qkqh','QDqh']:
                        #delete -1, and then add new consonent
                    self.deletePreviousChars(3,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3
                    # AUTO ADD HALANT
                    print("adding auto halant")
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')
                elif key_4 in ['chqh']:
                    #delete -1, and then add new consonent
                    self.deletePreviousChars(2,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3
                    # AUTO ADD HALANT
                    print("adding auto halant")
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')
                elif key_4 in ['JYqA']:
                    #delete -1, and then add new consonent
                    self.deletePreviousChars(3,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3

                else:
                    #delete -1, and then add new consonent
                    self.deletePreviousChars(3,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3
                    # AUTO ADD HALANT
                    print("adding auto halant")
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_3 in self.C:
                    print("in key_3")
                    #delete -1, and then add new consonent
                    # AUTO ADD HALANT
                    if key_3 in ['ZHA']: # no auto-halant after sindhi letters, and other rare consonents
                        self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.C[key_3])

                    else: 
                        self.deletePreviousChars(2,cursor)
                        self.insertPlainText(self.C[key_3])
                        print("adding auto halant")
                        self.insertPlainText(self.misc['q'])
                        self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                    self.makeNextVowelDependent = True
                    self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_3 in self.misc: #AUM
                print("in key_3")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.misc[key_3])
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True # every misc element resets this FLAG aai, taai vs m'ou'

            elif key_2 in self.C:
                print("in key_2")

                if key_2 in ['Zg','Zj','ZD','Zb','ZK']: # no auto-halant after sindhi letters, and other rare consonents
                        #self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.C[key_2])
                elif key_2 in ['ch','Qk','Qg','Qj','QP','QD','Qy','Qn','Qr','QL','Qv']: # no auto-halant after sindhi letters, and other rare consonents
                        #self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.C[key_2])
                        self.insertPlainText(self.misc['q'])
                        self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')  
                else:
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.C[key_2])
                    # AUTO ADD HALANT (if key_2[0] not 'q')
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_2 in self.misc:
                print("in key_2")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.misc[key_2])
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True # every misc element resets this FLAG aai, taai vs m'ou'

            elif key_1 in self.misc:
                print("in key_1")
                if key_2 == 'q ':
                    #if a space ' ' follows an ununsed halant, first delete the halant, and then add the space.
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.misc[' '])
                    self.makeNextVowelDependent = False
                    self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'
                elif key_1 in ['w','W']: #Inserting ZWNJ or ZWJ
                    self.insertPlainText(self.misc[key_1])
                    #self.makeNextVowelDependent = True a ZWNJ/ZWJ has to be followed by a consonent,or a space
                else:
                    #what kind of vowel should it be next              
                    self.insertPlainText(self.misc[key_1])
                    self.makeNextVowelDependent = False
                    self.use2CharsVowelNext = True # every misc element resets this FLAG aai, taai vs m'ou'

            # VOWEL LOGIC BEGINS
            # HARD CODED LOGIC
            elif key_5 in ['LqLqi','LqLqI']:
                print("in key_5")    
                self.deletePreviousChars(4,cursor)
                self.insertPlainText(self.V[key_5])
                self.makeNextVowelDependent = False

            elif key_5 in ['Lqlqi','LqlqI']:
                print("in key_5")    
                self.deletePreviousChars(4,cursor)
                self.insertPlainText(self.v[key_5])
                self.makeNextVowelDependent = False

            # for a given language, we want to check if key 3 is in repective english mappings of ृ, ऋ, ॠ
            elif self.current_script not in ['indus','manipuri'] and key_3 in [list(self.v.keys())[list(self.v.values()).index('ृ')],list(self.v.keys())[list(self.v.values()).index('\u0944')]]: 
                print("in key_3")
                self.deletePreviousChars(2,cursor)
                self.insertPlainText(self.v[key_3])
                self.makeNextVowelDependent = False    
                    
            elif self.current_script not in ['indus','manipuri'] and key_3 in [list(self.V.keys())[list(self.V.values()).index('ऋ')],list(self.V.keys())[list(self.V.values()).index('ॠ')]]:
                print("in key_3")
                #self.deletePreviousChars(2,cursor)
                self.insertPlainText(self.V[key_3])
                self.makeNextVowelDependent = False

            elif key_3 in ['zAU','zEE']:
                print("in key_3")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.V[key_3])
                self.makeNextVowelDependent = False

            elif key_3 in ['zau']:
                print("in key_3")
                self.deletePreviousChars(2,cursor)
                self.insertPlainText(self.v[key_3])
                self.makeNextVowelDependent = False

            elif key_2 in ['zO']:
                print("in key_2")
                #self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.V[key_2])
                self.makeNextVowelDependent = False

            elif key_2 in ['zo','ze']:
                print("in key_2")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.v[key_2])
                self.makeNextVowelDependent = False
    
            elif key_2 in ['AA','II','UU','EE','AU']: # alson includes rare characters "Curvy matras"
                print("in key_2")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.V[key_2])
                self.makeNextVowelDependent = False

            elif key_1 in ['A','I','U','E','O']:
                if key_2 not in ['AE','AO','aE','aO']:
                    print("in key_1")
                    self.insertPlainText(self.V[key_1])
                    self.makeNextVowelDependent = False
                else:
                    print('Writing - AO, AE, aE, aO')
                    if key_2[0] in ['a']:
                        if self.makeNextVowelDependent==False and key_1!='a':
                            self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.v[key_2])
                    if key_2[0] in ['A']:
                        self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.V[key_2])

            
            # CONTEXT BASED LOGIC
            elif (key_2 in self.V or key_2 in self.v) and self.use2CharsVowelNext:
                if not self.makeNextVowelDependent:
                    print(f"2 Chars used: inserting INDEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.V[key_2])
                    self.makeNextVowelDependent = False

                    #handing some edge case
                    if key_2 == 'aa':
                        self.use2CharsVowelNext = False

                elif key_2[0] == " ":
                    print(f"2 Chars used: inserting INDEPENDENT vowel because key_2[0] is ' '")
                    self.insertPlainText(self.V[key_2])
                    self.makeNextVowelDependent = True
                elif self.makeNextVowelDependent:
                    print(f"2 Chars used: inserting DEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")

                    #handing some edge cases
                    if key_2[1]!='a' and key_2[0] == key_2[1]: # for ii and uu, delete the previous dependent vowel and add next. But not for aa
                        self.deletePreviousChars(1,cursor)
                    elif key_2 == 'ou':
                        self.deletePreviousChars(1,cursor)
                    if key_2 == 'aa':
                        self.use2CharsVowelNext = False

                    self.makeNextVowelDependent = False
                    self.insertPlainText(self.v[key_2])     

            elif key_1 in self.V or key_1 in self.v:
                # POST AUTO HALANT LOGIC
                if key_2 == 'qa':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.makeNextVowelDependent = True 
                elif key_2 == 'qi':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['i'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'qo':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['o'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'qe':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['e'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'qu':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['u'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'q ':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.misc[' '])
                    self.makeNextVowelDependent = False 

                ## POST AUTO HALANT LOGIC ENDS
                    
  
                    
                elif not self.makeNextVowelDependent:
                    print(f"1 Char used: inserting INDEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")
                    self.insertPlainText(self.V[key_1])
                    self.makeNextVowelDependent = False
                    self.use2CharsVowelNext = True
                elif key_2[0] == " ":
                    print(f"1 Char used: used: inserting INDEPENDENT vowel because key_2[0] is ' '")
                    self.insertPlainText(self.V[key_1])
                    self.makeNextVowelDependent = True

                elif self.makeNextVowelDependent:
                    print(f"1 Char used: inserting DEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")
                    self.insertPlainText(self.v[key_1])
                    self.makeNextVowelDependent = True

            # SINGLE LETTER CONSONENT GETS LAST PRECIDENCE
            elif key_1 in self.C:
                if key_3[-1] =='r': #no auto halant
                        self.insertPlainText(self.C[key_1])
                else:
                    #adding consonent + auto halant
                    self.insertPlainText(self.C[key_1])
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'
            # INSERT UNKNOWN CHARS AS THEY ARE, or NOT DISPLAY THEM(q,R)
            else:
                print("unmapped key")
                print(event.key())
                if event.key() == 81: # Q
                    pass
                elif event.key() == 90: # Z, z
                    pass
                elif event.key() == 67: # c, x
                    pass
                elif event.key() == 82: # R
                    pass
                # R --> RRi,RRI,Rri,RrI
                # J --> JYA

                # elif key_1[0] == 'R':
                #     pass
                else:
                    #this is only for display, we are recording the key press in the backend anyways
                    self.insertPlainText(event.text())
                    print(event.key())
                    self.makeNextVowelDependent = False
            
            #self.label_trackScriptKeys.setText(str([l+'\n'for l in re.split(r'\n+', self.toPlainText().strip())]))
    
    def keyReleaseEvent(self, event):
        print(f"keyReleased: {event.text()}")
        if event.key() == QtCore.Qt.Key_QuoteLeft:  # Backtick key
            self.is_backtick_pressed = False
        return super(TextEdit, self).keyReleaseEvent(event)
                                                                                                                                                                                                                                                                                                                         
    # def onCursorPositionChangedByMouseOrArrow(self,pos):
    #   CLEAR label_trackAllKeys TEXT

    def deletePreviousChars(self, n, cursor):
        """
        Delete the previous n characters while maintaining text formatting.
        If there are fewer than n chars before the cursor, delete what's available.
        """

        # Store the current position
        current_position = cursor.position()
        # Calculate the start position for deletion
        start_position = max(1, current_position - n)

        if current_position-start_position<n:
            print('bug here!!')
            # Select the text to be deleted
            cursor.setPosition(start_position, QTextCursor.MoveAnchor)
            cursor.setPosition(current_position, QTextCursor.KeepAnchor)
            # Delete the selected text
            cursor.removeSelectedText()
            #todo replace character present from 0 to 1 with ZWNJ
            cursor.setPosition(0, QTextCursor.MoveAnchor)
            cursor.setPosition(1, QTextCursor.KeepAnchor)
            cursor.insertText('\u200C')  # ZWNJ (Zero-Width Non-Joiner)
        else:
            # Select the text to be deleted
            cursor.setPosition(start_position, QTextCursor.MoveAnchor)
            cursor.setPosition(current_position, QTextCursor.KeepAnchor)
            # Delete the selected text
            cursor.removeSelectedText()
        # Set the new cursor position
        self.setTextCursor(cursor)
        print(f"Deleted previous {min(n, current_position)} chars")



    def canInsertFromMimeData(self, source):
        
        if source.hasImage():
            return True
        else:
            return super(TextEdit, self).canInsertFromMimeData(source)

    def insertFromMimeData(self, source):
        
        cursor = self.textCursor()
        document = self.document()

        if source.hasUrls():

            for u in source.urls():
                file_ext = splitext(str(u.toLocalFile()))
                if u.isLocalFile() and file_ext in IMAGE_EXTENSIONS:
                    image = QImage(u.toLocalFile())
                    document.addResource(QTextDocument.ImageResource, u, image)
                    cursor.insertImage(u.toLocalFile())

                else:
                    # If we hit a non-image or non-local URL break the loop and fall out
                    # to the super call & let Qt handle it
                    break

            else:
                # If all were valid images, finish here.
                return
                


        elif source.hasImage():
            image = source.imageData()
            uuid = hexuuid()
            document.addResource(QTextDocument.ImageResource, uuid, image)
            cursor.insertImage(uuid)
            return

        super(TextEdit, self).insertFromMimeData(source)

    def setMarkdown(self, markdown: typing.Optional[str]) -> None:
            """Sets the content to the provided markdown string or clears it if None is provided."""
            if markdown is None:
                pass
                # self.content = ""
            else:
                print("this rannnnnnnnn")
                print("noneee")
                #self.content = markdown
                # self.content = 'hahha'
                # print(self.content)
                
    def setDocument(self, document):
            print(f"setDocument called. New document ID: {id(document)}")
            super().setDocument(document)

























class TextEdit_Indus(QTextEdit):
    def __init__(self,label_trackAllKeys, *args, **kwargs):
        super(self.__class__, self).__init__(*args, **kwargs)

        self.label_trackAllKeys = label_trackAllKeys
        self.current_script: str = ""
        # IMPORTING HINDI MAPPING DICTIONARIES
        from mappings.mappings_v1 import C,V,v,misc   #TODO change these to be indus by default
        self.C = C #consonents
        self.V = V #Independent Vowels
        self.v = v #dependent vowels
        self.misc = misc #nukta, halant, numbers and misc symbols
        self.english_bypass = False
        self.is_backtick_pressed = False

        print("using Indus text edit")

        # Variables used for Language rules
        self.lastChars= '   '#str
        self.makeNextVowelDependent = False
        self.use2CharsVowelNext = True
    
    def pasteWithoutFormatting(self):
        clipboard = QtWidgets.QApplication.clipboard()
        text = clipboard.text()
        self.insertPlainText(text)
        font = self.document().defaultFont()
        self.document().setDefaultFont(font)


    def keyPressEvent(self, event):
        print("_______________________________")
        print("use2CharVowelNext = "+str(self.use2CharsVowelNext))
        print(f"keyPressed: {event.text()}")
        print(f"keyPressed_id: {event}")
        # Process the default keys (spacebar, backspace, return/enter) like how we would normally
        if event.key() in [QtCore.Qt.Key_Enter,QtCore.Qt.Key_Return]:
            self.makeNextVowelDependent = False # new line should start with an independent vowele
            self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'\n')
            #self.label_trackScriptKeys.setText(self.label_trackScriptKeys.text()+"\n")
            return super(TextEdit, self).keyPressEvent(event) 
        elif event.key() in [QtCore.Qt.Key_Left,QtCore.Qt.Key_Right]: #, QtCore.Qt.AltModifier
            return super(TextEdit, self).keyPressEvent(event) 
        elif event.key() in [QtCore.Qt.Key_Shift]:
            return super(TextEdit, self).keyPressEvent(event)
        elif event.key() in [QtCore.Qt.Key_F5]: # refresh all flags and clear context
                self.label_trackAllKeys.setText('   ')#str
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True
                print("F5 pressed")
                return super(TextEdit, self).keyPressEvent(event)
        # if int(event.modifiers()) == (QtCore.Qt.ControlModifier+QtCore.Qt.AltModifier):
        #         # Alt Gr or Ctrl+Alt pressed 
        #         print('Alt Gr or Ctrl+Alt pressed')
        #         #event.ignore()
        #         # THE FOLLOWING IS DEDICATED LOGIC FOR DEPENDENT SINDHI VOWELS āū,ōū,ō,ē
        #         # FOR THE INDEPENDENT VOWELS, SHIFT KEY IS ALSO PRESSED, HENCE THIS SECTION DOES NOT GET EXECUTED THEN.
                
        #         # TODO if key pressed is ōūā, then insert text(in display and in context), 
                # without that wierd alt functionality.

        if event.modifiers() & QtCore.Qt.ControlModifier:
            if event.key() == QtCore.Qt.Key_Z:
                self.undo()
            elif event.key() == QtCore.Qt.Key_C:
                self.copy()
            elif event.key() == QtCore.Qt.Key_V:
                self.pasteWithoutFormatting()
            elif event.key() == QtCore.Qt.Key_X:
                self.cut()
        elif event.modifiers() & QtCore.Qt.AltModifier:
            event.ignore()
            print('ALT PRESSED')
        elif event.key() == QtCore.Qt.Key_Backspace:
            if len(self.toPlainText()) >1: 
                # do backspace in the keypress tracker also..
                self.label_trackAllKeys.setText(self.label_trackAllKeys.text()[:-1])
                return super(TextEdit, self).keyPressEvent(event)
            else:
                print("CLEARED")
                self.clear()
        elif event.key() == QtCore.Qt.Key_Down:
            cursor = self.textCursor()
            cursor.movePosition(cursor.Start)
            self.setTextCursor(cursor)
            return super(TextEdit, self).keyPressEvent(event)
        elif event.key() == QtCore.Qt.Key_Up:
            cursor = self.textCursor()
            cursor.movePosition(cursor.End)
            self.setTextCursor(cursor)
            return super(TextEdit, self).keyPressEvent(event)
        

        elif self.english_bypass == True:
            print("in here typing ENLGISH")
            self.insertPlainText(event.text())
        elif self.is_backtick_pressed: #################
            if QtCore.Qt.Key_A <= event.key() <= QtCore.Qt.Key_Z:
                print(f"BACKTICK + {event.text().upper()} PRESSED")
                print("in here typing ENLGISH")
                self.insertPlainText(event.text())
                self.label_trackAllKeys.setText('   ')#str
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True
        elif event.key() == QtCore.Qt.Key_QuoteLeft:  # Backtick key
            self.is_backtick_pressed = True
        else:
            self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+event.text())
            self.lastChars = self.label_trackAllKeys.text()[-5:]

            cursor = self.textCursor()

            key_5 = ''.join(self.lastChars[-5:])
            key_4 = ''.join(self.lastChars[-4:])
            key_3 = ''.join(self.lastChars[-3:])
            key_2 = ''.join(self.lastChars[-2:])
            key_1 = ''.join(self.lastChars[-1:])

            # Handling errors at the start of the code
            while len(key_5)<5:
                key_5 = ' '+ key_5
            while len(key_4)<4:
                key_4 = ' '+ key_4
            while len(key_3)<3:
                key_3 = ' '+ key_3
            while len(key_2)<2:
                key_2 = ' '+ key_2

            print("5 keys are:")
            print(key_5)
            print(key_4)
            print(key_3)
            print(key_2)
            print(key_1)
        
            #EDGE CASES
            if key_5[0] in self.C and key_5[2] == 'r' and key_5[4] == 'u':
                    self.deletePreviousChars(4 ,cursor)
                    self.insertPlainText(self.C[key_5[0]]+self.v['Rri'])
        
            #CONSONENTS AND MISC
            #detect matching keys in the last 4,3,2,1 chars
            #independent when last is " ", or last is Dependent
            #Dependent when last is consonent, or last is in a,e,i,o,u
            elif key_5 in self.C:
                #delete -1, and then add new consonent
                print("in key_5")
                if key_5 in ['sqhqr']:
                    self.deletePreviousChars(2 ,cursor)
                    self.insertPlainText(self.C[key_5])
                    
                else:
                    self.deletePreviousChars(4 ,cursor)
                    self.insertPlainText(self.C[key_5])
                    if key_5 not in ['DqDqA']:
                        # AUTO ADD HALANT
                        self.insertPlainText(self.misc['q'])
                        self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q') #logging in english

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_4 in self.C:
                print("in key_4")
                if key_4 in ['Qkqh','QDqh']:
                        #delete -1, and then add new consonent
                    self.deletePreviousChars(3,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3
                    # AUTO ADD HALANT
                    print("adding auto halant")
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')
                elif key_4 in ['chqh']:
                    #delete -1, and then add new consonent
                    self.deletePreviousChars(2,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3
                    # AUTO ADD HALANT
                    print("adding auto halant")
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')
                elif key_4 in ['JYqA']:
                    #delete -1, and then add new consonent
                    self.deletePreviousChars(3,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3

                else:
                    #delete -1, and then add new consonent
                    self.deletePreviousChars(3,cursor) #2
                    self.insertPlainText(self.C[key_4]) #key_3
                    # AUTO ADD HALANT
                    print("adding auto halant")
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_3 in self.C:
                    print("in key_3")
                    #delete -1, and then add new consonent
                    # AUTO ADD HALANT
                    if key_3 in ['ZHA']: # no auto-halant after sindhi letters, and other rare consonents
                        self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.C[key_3])

                    else: 
                        self.deletePreviousChars(2,cursor)
                        self.insertPlainText(self.C[key_3])
                        print("adding auto halant")
                        self.insertPlainText(self.misc['q'])
                        self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                    self.makeNextVowelDependent = True
                    self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_3 in self.misc: #AUM
                print("in key_3")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.misc[key_3])
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True # every misc element resets this FLAG aai, taai vs m'ou'

            elif key_2 in self.C:
                print("in key_2")

                if key_2 in ['Zg','Zj','ZD','Zb','ZK']: # no auto-halant after sindhi letters, and other rare consonents
                        #self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.C[key_2])
                elif key_2 in ['ch','Qk','Qg','Qj','QP','QD','Qy','Qn','Qr','QL','Qv']: # no auto-halant after sindhi letters, and other rare consonents
                        #self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.C[key_2])
                        self.insertPlainText(self.misc['q'])
                        self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')  
                else:
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.C[key_2])
                    # AUTO ADD HALANT (if key_2[0] not 'q')
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'

            elif key_2 in self.misc:
                print("in key_2")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.misc[key_2])
                self.makeNextVowelDependent = False
                self.use2CharsVowelNext = True # every misc element resets this FLAG aai, taai vs m'ou'

            elif key_1 in self.misc:
                print("in key_1")
                if key_2 == 'q ':
                    #if a space ' ' follows an ununsed halant, first delete the halant, and then add the space.
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.misc[' '])
                    self.makeNextVowelDependent = False
                    self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'
                elif key_1 in ['w','W']: #Inserting ZWNJ or ZWJ
                    self.insertPlainText(self.misc[key_1])
                    #self.makeNextVowelDependent = True a ZWNJ/ZWJ has to be followed by a consonent,or a space
                else:
                    #what kind of vowel should it be next              
                    self.insertPlainText(self.misc[key_1])
                    self.makeNextVowelDependent = False
                    self.use2CharsVowelNext = True # every misc element resets this FLAG aai, taai vs m'ou'

            # VOWEL LOGIC BEGINS
            # HARD CODED LOGIC
            elif key_5 in ['LqLqi','LqLqI']:
                print("in key_5")    
                self.deletePreviousChars(4,cursor)
                self.insertPlainText(self.V[key_5])
                self.makeNextVowelDependent = False

            elif key_5 in ['Lqlqi','LqlqI']:
                print("in key_5")    
                self.deletePreviousChars(4,cursor)
                self.insertPlainText(self.v[key_5])
                self.makeNextVowelDependent = False

            # for a given language, we want to check if key 3 is in repective english mappings of ृ, ऋ, ॠ
            elif self.current_script not in ['indus','manipuri'] and key_3 in [list(self.v.keys())[list(self.v.values()).index('ृ')],list(self.v.keys())[list(self.v.values()).index('\u0944')]]: 
                print("in key_3")
                self.deletePreviousChars(2,cursor)
                self.insertPlainText(self.v[key_3])
                self.makeNextVowelDependent = False    
                    
            elif self.current_script not in ['indus','manipuri'] and key_3 in [list(self.V.keys())[list(self.V.values()).index('ऋ')],list(self.V.keys())[list(self.V.values()).index('ॠ')]]:
                print("in key_3")
                #self.deletePreviousChars(2,cursor)
                self.insertPlainText(self.V[key_3])
                self.makeNextVowelDependent = False

            elif key_3 in ['zAU','zEE']:
                print("in key_3")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.V[key_3])
                self.makeNextVowelDependent = False

            elif key_3 in ['zau']:
                print("in key_3")
                self.deletePreviousChars(2,cursor)
                self.insertPlainText(self.v[key_3])
                self.makeNextVowelDependent = False

            elif key_2 in ['zO']:
                print("in key_2")
                #self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.V[key_2])
                self.makeNextVowelDependent = False

            elif key_2 in ['zo','ze']:
                print("in key_2")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.v[key_2])
                self.makeNextVowelDependent = False
    
            elif key_2 in ['AA','II','UU','EE','AU']: # alson includes rare characters "Curvy matras"
                print("in key_2")
                self.deletePreviousChars(1,cursor)
                self.insertPlainText(self.V[key_2])
                self.makeNextVowelDependent = False

            elif key_1 in ['A','I','U','E','O']:
                if key_2 not in ['AE','AO','aE','aO']:
                    print("in key_1")
                    self.insertPlainText(self.V[key_1])
                    self.makeNextVowelDependent = False
                else:
                    print('Writing - AO, AE, aE, aO')
                    if key_2[0] in ['a']:
                        if self.makeNextVowelDependent==False and key_1!='a':
                            self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.v[key_2])
                    if key_2[0] in ['A']:
                        self.deletePreviousChars(1,cursor)
                        self.insertPlainText(self.V[key_2])

            
            # CONTEXT BASED LOGIC
            elif (key_2 in self.V or key_2 in self.v) and self.use2CharsVowelNext:
                if not self.makeNextVowelDependent:
                    print(f"2 Chars used: inserting INDEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.V[key_2])
                    self.makeNextVowelDependent = False

                    #handing some edge case
                    if key_2 == 'aa':
                        self.use2CharsVowelNext = False

                elif key_2[0] == " ":
                    print(f"2 Chars used: inserting INDEPENDENT vowel because key_2[0] is ' '")
                    self.insertPlainText(self.V[key_2])
                    self.makeNextVowelDependent = True
                elif self.makeNextVowelDependent:
                    print(f"2 Chars used: inserting DEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")

                    #handing some edge cases
                    if key_2[1]!='a' and key_2[0] == key_2[1]: # for ii and uu, delete the previous dependent vowel and add next. But not for aa
                        self.deletePreviousChars(1,cursor)
                    elif key_2 == 'ou':
                        self.deletePreviousChars(1,cursor)
                    if key_2 == 'aa':
                        self.use2CharsVowelNext = False

                    self.makeNextVowelDependent = False
                    self.insertPlainText(self.v[key_2])     

            elif key_1 in self.V or key_1 in self.v:
                # POST AUTO HALANT LOGIC
                if key_2 == 'qa':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.makeNextVowelDependent = True 
                elif key_2 == 'qi':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['i'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'qo':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['o'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'qe':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['e'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'qu':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.v['u'])
                    self.makeNextVowelDependent = True
                elif key_2 == 'q ':
                    print('auto deleting halant..')
                    self.deletePreviousChars(1,cursor)
                    self.insertPlainText(self.misc[' '])
                    self.makeNextVowelDependent = False 

                ## POST AUTO HALANT LOGIC ENDS
                    
  
                    
                elif not self.makeNextVowelDependent:
                    print(f"1 Char used: inserting INDEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")
                    self.insertPlainText(self.V[key_1])
                    self.makeNextVowelDependent = False
                    self.use2CharsVowelNext = True
                elif key_2[0] == " ":
                    print(f"1 Char used: used: inserting INDEPENDENT vowel because key_2[0] is ' '")
                    self.insertPlainText(self.V[key_1])
                    self.makeNextVowelDependent = True

                elif self.makeNextVowelDependent:
                    print(f"1 Char used: inserting DEPENDENT vowel because self.makeNextVowelDependent is {self.makeNextVowelDependent}")
                    self.insertPlainText(self.v[key_1])
                    self.makeNextVowelDependent = True

            # SINGLE LETTER CONSONENT GETS LAST PRECIDENCE
            elif key_1 in self.C:
                if key_3[-1] =='r': #no auto halant
                        self.insertPlainText(self.C[key_1])
                else:
                    #adding consonent + auto halant
                    self.insertPlainText(self.C[key_1])
                    self.insertPlainText(self.misc['q'])
                    self.label_trackAllKeys.setText(self.label_trackAllKeys.text()+'q')

                self.makeNextVowelDependent = True
                self.use2CharsVowelNext = True # every consonent resets this FLAG aai, taai vs m'ou'
            # INSERT UNKNOWN CHARS AS THEY ARE, or NOT DISPLAY THEM(q,R)
            else:
                print("unmapped key")
                print(event.key())
                if event.key() == 81: # Q
                    pass
                elif event.key() == 90: # Z, z
                    pass
                elif event.key() == 67: # c, x
                    pass
                elif event.key() == 82: # R
                    pass
                # R --> RRi,RRI,Rri,RrI
                # J --> JYA

                # elif key_1[0] == 'R':
                #     pass
                else:
                    #this is only for display, we are recording the key press in the backend anyways
                    self.insertPlainText(event.text())
                    print(event.key())
                    self.makeNextVowelDependent = False
            
            #self.label_trackScriptKeys.setText(str([l+'\n'for l in re.split(r'\n+', self.toPlainText().strip())]))
    
    def keyReleaseEvent(self, event):
        print(f"keyReleased: {event.text()}")
        if event.key() == QtCore.Qt.Key_QuoteLeft:  # Backtick key
            self.is_backtick_pressed = False
        return super(TextEdit_Indus, self).keyReleaseEvent(event)
                                                                                                                                                                                                                                                                                                                         
    # def onCursorPositionChangedByMouseOrArrow(self,pos):
    #   CLEAR label_trackAllKeys TEXT

    def deletePreviousChars(self, n, cursor):
        """
        Delete the previous n characters while maintaining text formatting.
        If there are fewer than n chars before the cursor, delete what's available.
        """

        # Store the current position
        current_position = cursor.position()
        # Calculate the start position for deletion
        start_position = max(1, current_position - n)

        if current_position-start_position<n:
            print('bug here!!')
            # Select the text to be deleted
            cursor.setPosition(start_position, QTextCursor.MoveAnchor)
            cursor.setPosition(current_position, QTextCursor.KeepAnchor)
            # Delete the selected text
            cursor.removeSelectedText()
            #todo replace character present from 0 to 1 with ZWNJ
            cursor.setPosition(0, QTextCursor.MoveAnchor)
            cursor.setPosition(1, QTextCursor.KeepAnchor)
            cursor.insertText('\u200C')  # ZWNJ (Zero-Width Non-Joiner)
        else:
            # Select the text to be deleted
            cursor.setPosition(start_position, QTextCursor.MoveAnchor)
            cursor.setPosition(current_position, QTextCursor.KeepAnchor)
            # Delete the selected text
            cursor.removeSelectedText()
        # Set the new cursor position
        self.setTextCursor(cursor)
        print(f"Deleted previous {min(n, current_position)} chars")



    def canInsertFromMimeData(self, source):
        
        if source.hasImage():
            return True
        else:
            return super(TextEdit, self).canInsertFromMimeData(source)

    def insertFromMimeData(self, source):
        
        cursor = self.textCursor()
        document = self.document()

        if source.hasUrls():

            for u in source.urls():
                file_ext = splitext(str(u.toLocalFile()))
                if u.isLocalFile() and file_ext in IMAGE_EXTENSIONS:
                    image = QImage(u.toLocalFile())
                    document.addResource(QTextDocument.ImageResource, u, image)
                    cursor.insertImage(u.toLocalFile())

                else:
                    # If we hit a non-image or non-local URL break the loop and fall out
                    # to the super call & let Qt handle it
                    break

            else:
                # If all were valid images, finish here.
                return
                


        elif source.hasImage():
            image = source.imageData()
            uuid = hexuuid()
            document.addResource(QTextDocument.ImageResource, uuid, image)
            cursor.insertImage(uuid)
            return

        super(TextEdit, self).insertFromMimeData(source)

    def setMarkdown(self, markdown: typing.Optional[str]) -> None:
            """Sets the content to the provided markdown string or clears it if None is provided."""
            if markdown is None:
                pass
                # self.content = ""
            else:
                print("this rannnnnnnnn")
                print("noneee")
                #self.content = markdown
                # self.content = 'hahha'
                # print(self.content)
                
    def setDocument(self, document):
            print(f"setDocument called. New document ID: {id(document)}")
            super().setDocument(document)